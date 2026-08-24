"""Extracts plain text from uploaded resume files (PDF / DOCX / TXT)."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import docx
import pdfplumber

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class ResumeParsingError(Exception):
    pass


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ResumeParsingError(
            f"Unsupported file type '{ext}'. Please upload a PDF, DOCX, or TXT resume."
        )

    try:
        if ext == ".pdf":
            return _extract_pdf(file_bytes)
        if ext == ".docx":
            return _extract_docx(file_bytes)
        return file_bytes.decode("utf-8", errors="ignore")
    except ResumeParsingError:
        raise
    except Exception as exc:  # pragma: no cover - defensive
        raise ResumeParsingError(f"Failed to parse resume: {exc}") from exc


def _extract_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    text = "\n".join(text_parts).strip()
    if not text:
        raise ResumeParsingError(
            "Could not extract any text from this PDF. It may be a scanned image without a text layer."
        )
    return text


def _extract_docx(file_bytes: bytes) -> str:
    document = docx.Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n".join(paragraphs).strip()
    if not text:
        raise ResumeParsingError("Could not extract any text from this DOCX file.")
    return text
